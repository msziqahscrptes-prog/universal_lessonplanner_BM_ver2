import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from io import BytesIO

# --- 1. KONFIGURASI HALAMAN ---
st.set_page_config(page_title="Perancang Master Universal", layout="wide")
st.title("🎓 PTES Penjana Rancangan Mengajar Universal (BM)")

# --- INPUT KUNCI API (DI BAHAGIAN ATAS) ---
user_api_key = st.text_input(
    "🔑 Masukkan Kunci API Gemini Anda:", 
    type="password", 
    help="Dapatkan kunci API anda dari Google AI Studio menggunakan akaun Gmail anda."
)

# Fungsi untuk memeriksa dan memuatkan model secara dinamik
def get_working_model(api_key):
    try:
        genai.configure(api_key=api_key)
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except Exception as e:
        st.error(f"Ralat Kunci API atau Masalah Sambungan: {str(e)}")
        return None
    return "models/gemini-1.5-flash"  # Sandaran Lalai


# Memproses pengesahan model
selected_model_name = None
if user_api_key:
    selected_model_name = get_working_model(user_api_key)
    if selected_model_name:
        st.info(f"Sistem disambungkan. Model Aktif: {selected_model_name}")
else:
    st.warning("⚠️ Sila masukkan Kunci API Gemini peribadi anda di atas untuk bermula.")


# --- 2. LOGIK AI (BAHASA MELAYU DENGAN SENARAI NOMBOR & HURUF BESAR) ---
def generate_advanced_plan_bm(topic, syllabus, extra_context, api_key, model_name):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    
    # Arahan spesifik untuk huruf besar, tanpa asteris, dan senarai bernombor
    prompt = f"""
    Topik: {topic}. Kod Silibus: {syllabus}. Konteks Tambahan: {extra_context}.
    Janakan rancangan mengajar profesional dalam Bahasa Melayu sepenuhnya.
    JANGAN gunakan simbol dwi-asterisk (**) sama sekali dalam teks output.
    Gunakan senarai bernombor (1, 2, 3, 4) dan bukannya senarai peluru (bullet points) untuk setiap poin huraian criteria.
    
    Gunakan penanda (markers) dalam HURUF BESAR yang tepat ini untuk struktur dokumen:
    
    SECTION: TOPIK PELAJARAN
    [Paparkan topik input di sini]
    
    SECTION: OBJEKTIF PELAJARAN
    [Sediakan 4 mata dalam bentuk senarai bernombor 1 hingga 4]
    
    SECTION: HASIL PELAJARAN
    [Sediakan 4 mata dalam bentuk senarai bernombor 1 hingga 4]
    
    SECTION: KRITERIA KEJAYAAN
    [Sediakan 4 mata dalam bentuk senarai bernombor 1 hingga 4]
    
    SECTION: PRASYARAT
    [1 mata utama]
    
    SECTION: KATA KUNCI
    [Sediakan 6 item kata kunci pelajaran bernombor 1 hingga 6]
    
    SECTION: KBAT
    [Sediakan 4 domain utama dari Taksonomi Bloom bernombor 1 hingga 4]
    
    SECTION: KEWARGANEGARAAN DIGITAL
    [Sediakan 4 mata bernombor 1 hingga 4 mengenai penggunaan teknologi beretika/Chromebook/Canva/YouTube]

    SECTION: KANDUNGAN PEMBUKAAN PELAJARAN
    [Aktiviti set induksi/pemicu minda dan pelan transisi]

    SECTION: STRATEGI DIFERENSIASI (HIJAU)
    - HA (Murid Pencapaian Tinggi): [1 aktiviti mencabar dalam bentuk senarai bernombor]

    SECTION: STRATEGI DIFERENSIASI (KUNING)
    - MA (Murid Pencapaian Sederhana): [1 aktiviti teras dalam bentuk senarai bernombor]

    SECTION: STRATEGI DIFERENSIASI (MERAH)
    - LA (Murid Pencapaian Rendah): [1 aktiviti sokongan/berperancah dalam bentuk senarai bernombor]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN SATU (15 MINIT)
    - Aktiviti 1: [Huraian]
    - ----------------------------------------------------------------------------
    - Persediaan Guru: [Langkah demi langkah sebelum kelas bermula]
    - ----------------------------------------------------------------------------
    - Objektif: [3 mata utama menggunakan senarai bernombor]
    - ----------------------------------------------------------------------------
    - Tugasan Murid: [Butiran langkah demi langkah murid]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN DUA (15 MINIT)
    - Aktiviti 2: [Huraian]
    - -----------------------------------------------------------------------------
    - Persediaan Guru: [Langkah demi langkah sebelum kelas bermula]
    - -----------------------------------------------------------------------------
    - Objektif: [3 mata utama menggunakan senarai bernombor]
    - -----------------------------------------------------------------------------
    - Tugasan Murid: [Butiran langkah demi langkah murid]
    
    SECTION: PLENARI (TIKET KELUAR)
    [Aktiviti penutup kelas sekitar 2-3 minit]

    SECTION: KERJA RUMAH
    [Tugasan yang diberikan berdasarkan topik]

    SECTION: CADANGAN TUGASAN SETERUSNYA
    [Aktiviti pemicu minda dan pelan transisi untuk sesi pengajaran hari esok]
    """
    try:
        response = model.generate_content(prompt)
        
        if response.candidates and response.candidates[0].content.parts:
            # Menggantikan baki dwi-asterisk yang mungkin terjana secara tidak sengaja
            cleaned_text = response.text.replace("**", "")
            return cleaned_text
        else:
            return "⚠️ AI mengembalikan respons kosong. Anda mungkin menekan butang terlalu cepat sehingga melebihi had kuota seminit (15 permintaan/minit). Sila tunggu 60 saat dan cuba lagi."
            
    except Exception as e:
        return f"Ralat Sistem: {str(e)}"

# --- 3. LOGIK EKSPORT WORD (HURUF BESAR) ---
def create_word_export(topic, syllabus, text):
    doc = Document()
    # Tajuk utama dokumen diubah ke format Huruf Besar
    doc.add_heading(f'PTES RANCANGAN MENGAJAR UNIVERSAL: {topic.upper()}', 0)

    # Jadual Pentadbiran Atas (Admin Header)
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Minggu No:", "Tarikh:"], ["Bilangan Murid:", "Hari:"], ["Tempat / Makmal:", "Durasi Kelas:"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
    doc.add_paragraph()

    # Penguraian Kandungan & Pembentukan Jadual Berkotak
    sections = text.split('SECTION:')
 
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        title = lines[0].strip().upper() # Memastikan tajuk diletakkan dalam HURUF BESAR penuh
        
        content = "\n".join(lines[1:]).strip().replace("**", "") 
        
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        table.cell(0, 0).text = content
        doc.add_paragraph()
     
    # Pengesahan Ketua Jabatan / HKP
    doc.add_page_break()
    doc.add_heading("KELULUSAN & ULASAN KETUA JABATAN / HKP", level=1)
    hod_table = doc.add_table(rows=2, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Ulasan / Catatan:"
    hod_table.rows[1].height = Pt(50)
    hod_table.cell(1, 0).text = "Tarikh:"; hod_table.cell(1, 1).text = "Tandatangan:"

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 4. GUI UTAMA ---
st.write("---")
st.info("Sila masukkan topik pelajaran, kod silibus subjek dan konteks tambahan seperti canva, youtube, atau infografik.")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Topik Pelajaran:")
with c2: u_syllabus = st.text_input("Kod Silibus:")
u_extra = st.text_area("Konteks Khas / Kata Kunci Tambahan (Opsional):")

if st.button("🚀 JANAKAN LANKAH PENGAJARAN UNIVERSAL"):
    if not user_api_key:
        st.error("❌ Sila masukkan Kunci API Gemini anda di bahagian atas sebelum menjana.")
    elif not u_topic or not u_syllabus:
        st.error("❌ Sila isi maklumat bidang Topik Pelajaran dan Kod Silibus.")
    else:
        with st.spinner("AI sedang mengintegrasikan kriteria ke dalam rancangan mengajar anda..."):
            result = generate_advanced_plan_bm(u_topic, u_syllabus, u_extra, user_api_key, selected_model_name)
            st.session_state['adv_plan_out_bm'] = result

if 'adv_plan_out_bm' in st.session_state:
    st.divider()
    st.subheader("Pratonton Draf AI")
    st.text_area("Kandungan", st.session_state['adv_plan_out_bm'], height=400)
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['adv_plan_out_bm'])
    st.download_button("📥 Muat Turun Versi Word (.docx)", doc_file, f"Universal_RM_{u_topic}.docx")

st.markdown("---")
st.caption("Rancangan Mengajar Universal ver.2.0 (BM) | Pencipta: Hjh Nurul Haziqah Hj Nordin | © 2026")
